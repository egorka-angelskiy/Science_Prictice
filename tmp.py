from docx import Document

document = open('Poyasnitelnaya_zapiska_Pleschev_Danil_021-1.docx', 'rb')
document = Document(document)


for i in range(len(document.paragraphs)):
	# print(i)
	
	if len(document.paragraphs[i].runs) > 0:
		
		for j in range(len(document.paragraphs[i].runs)):
			
			# print(i.runs[j].text == 'Таблица 1.1 – Функции меломана ')
			
			if document.paragraphs[i].runs[j].text == 'Таблица 1.1 – Функции меломана ':
				print(document.paragraphs[i].runs[j].text)

	# print('\n\n')

table_text = ''
for k in range(len(document.tables[0].rows)):
	for l in range(len(document.tables[0].rows[k].cells)):
		table_text += f'{document.tables[0].rows[k].cells[l].text:^60}'
	table_text += '\n'

print(table_text)
